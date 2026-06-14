# -*- coding: utf-8 -*-
"""
生成随机测试文章并转换为郑州大学工学报格式。
所有格式参数严格对齐模板文档实测值。
"""

from docx import Document
from docx.shared import Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import re
import subprocess
import tempfile
import copy

# ============================================================
# 格式常量（模板实测值）
# ============================================================
# 字号（run 覆盖值，EMU）
SIZE_TITLE = 228600         # 18pt
SIZE_AUTHOR = 177800        # 14pt
SIZE_AFFILIATION = 114300   # 9pt
SIZE_HEADING1 = 177800      # 14pt
SIZE_HEADING2 = 133350      # 10.5pt
SIZE_BODY = None            # 继承 Normal 样式 (sz=21, 10.5pt)
SIZE_ABSTRACT = 114300      # 9pt
SIZE_CAPTION = 114300       # 9pt
SIZE_REF = 114300           # 9pt
SIZE_ENG_TITLE = 152400     # 12pt
SIZE_ENG_AUTHOR = 139700    # 11pt（模板中英文作者部分 run 用此值）

# 缩进（EMU）
# 摘要：left=351twips, right=390twips
IND_ABSTRACT_LEFT = 351 * 635    # 222885
IND_ABSTRACT_RIGHT = 390 * 635   # 247650
# 正文：firstLine=406twips
IND_BODY_FIRST = 406 * 635       # 257810
# 参考文献：left=282twips（中文），left=375twips（英文）
IND_REF_ZH = 282 * 635           # 179070
IND_REF_EN = 375 * 635           # 238125

# 页面（twips）
PAGE_W = 11906
PAGE_H = 16838
MARGIN_TB = 1440
MARGIN_LR = 1800
COL_SPACE = 425

# Heading 样式段间距（用户指定值）
H1_SPACING_BEFORE = 340   # twips
H1_SPACING_AFTER = 330    # twips
H1_SPACING_LINE = 578
H1_SPACING_LINE_RULE = 'auto'

H2_SPACING_BEFORE = 260
H2_SPACING_AFTER = 260
H2_SPACING_LINE = 416
H2_SPACING_LINE_RULE = 'auto'

# 作者行间距
AUTHOR_SPACING_BEFORE = 157  # twips
AFFILIATION_SPACING_AFTER = 157  # twips


# ============================================================
# 工具函数
# ============================================================
def set_font(run, ea='宋体', fn='Times New Roman', size=None, bold=False):
    run.font.name = fn
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    if size:
        run.font.size = Emu(size)
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), ea)
    rFonts.set(qn('w:ascii'), fn)
    rFonts.set(qn('w:hAnsi'), fn)


def set_spacing(para, before=None, after=None, line=None, rule=None):
    pPr = para._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        para._element.insert(0, pPr)
    sp = pPr.find(qn('w:spacing'))
    if sp is None:
        sp = OxmlElement('w:spacing')
        pPr.append(sp)
    if before is not None:
        sp.set(qn('w:before'), str(before))
    if after is not None:
        sp.set(qn('w:after'), str(after))
    if line is not None:
        sp.set(qn('w:line'), str(line))
    if rule is not None:
        sp.set(qn('w:lineRule'), rule)


def set_indent(para, first=None, left=None, right=None):
    pPr = para._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        para._element.insert(0, pPr)
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind')
        pPr.append(ind)
    if first is not None:
        ind.set(qn('w:firstLine'), str(int(first / 635)))
    if left is not None:
        ind.set(qn('w:left'), str(int(left / 635)))
    if right is not None:
        ind.set(qn('w:right'), str(int(right / 635)))


def set_continuous(section):
    sectPr = section._sectPr
    type_el = sectPr.find(qn('w:type'))
    if type_el is None:
        type_el = OxmlElement('w:type')
        pgSz = sectPr.find(qn('w:pgSz'))
        if pgSz is not None:
            pgSz.addprevious(type_el)
        else:
            sectPr.append(type_el)
    type_el.set(qn('w:val'), 'continuous')


def setup_section(section, cols=1, col_space=COL_SPACE):
    sectPr = section._sectPr
    pgSz = sectPr.find(qn('w:pgSz'))
    if pgSz is None:
        pgSz = OxmlElement('w:pgSz')
        sectPr.append(pgSz)
    pgSz.set(qn('w:w'), str(PAGE_W))
    pgSz.set(qn('w:h'), str(PAGE_H))
    pgMar = sectPr.find(qn('w:pgMar'))
    if pgMar is None:
        pgMar = OxmlElement('w:pgMar')
        sectPr.append(pgMar)
    pgMar.set(qn('w:top'), str(MARGIN_TB))
    pgMar.set(qn('w:bottom'), str(MARGIN_TB))
    pgMar.set(qn('w:left'), str(MARGIN_LR))
    pgMar.set(qn('w:right'), str(MARGIN_LR))
    old_cols = sectPr.findall(qn('w:cols'))
    for c in old_cols:
        sectPr.remove(c)
    cols_elem = OxmlElement('w:cols')
    if cols > 1:
        cols_elem.set(qn('w:num'), str(cols))
        cols_elem.set(qn('w:space'), str(col_space))
    sectPr.append(cols_elem)


def make_heading_1(doc, text):
    """一级标题：14pt，段前段后0.5行(120twips)，行距240/auto"""
    p = doc.add_paragraph()
    p.style = doc.styles['Heading 1']
    m = re.match(r'^(\d+\s+)(.+)$', text)
    if m:
        run1 = p.add_run(m.group(1))
        set_font(run1, '宋体', 'Times New Roman', SIZE_HEADING1)
        run2 = p.add_run(m.group(2))
        set_font(run2, '黑体', 'Times New Roman', SIZE_HEADING1)
    else:
        run = p.add_run(text)
        set_font(run, '黑体', 'Times New Roman', SIZE_HEADING1)
    # 段前段后0.5行 = 120twips
    set_spacing(p, before=120, after=120, line=240, rule='auto')
    return p


def make_heading_2(doc, text):
    """二级标题：10.5pt，样式段间距 before=260 after=260 line=416/auto"""
    p = doc.add_paragraph()
    p.style = doc.styles['Heading 2']
    run = p.add_run(text)
    set_font(run, '黑体', 'Times New Roman', SIZE_HEADING2)
    # 段落级别覆盖：before=0, after=0, line=240/auto（单倍行距）
    set_spacing(p, before=0, after=0, line=240, rule='auto')
    return p


def make_body(doc, text):
    """正文：继承 Normal 样式 (10.5pt)，首行缩进 406twips，段前段后0"""
    p = doc.add_paragraph()
    set_indent(p, first=IND_BODY_FIRST)
    set_spacing(p, before=0, after=0)  # 段前段后必须为0磅
    run = p.add_run(text)
    # 不设置 size，继承 Normal 样式 sz=21 (10.5pt)
    set_font(run, '宋体', 'Times New Roman')
    return p


def make_caption(doc, text):
    """图表标题：9pt 黑体加粗，居中，行距 240/12pt exact"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, line=240, rule='exact')
    run = p.add_run(text)
    set_font(run, '黑体', 'Times New Roman', SIZE_CAPTION, bold=True)
    return p


def make_ref_zh(doc, text):
    """中文参考文献：9pt，left=282twips"""
    p = doc.add_paragraph()
    set_indent(p, left=IND_REF_ZH)
    run = p.add_run(text)
    set_font(run, '宋体', 'Times New Roman', SIZE_REF)
    return p


def make_ref_en(doc, text):
    """英文参考文献：9pt，left=375twips"""
    p = doc.add_paragraph()
    set_indent(p, left=IND_REF_EN)
    run = p.add_run(text)
    set_font(run, 'Times New Roman', 'Times New Roman', SIZE_REF)
    return p


def insert_formula(paragraph, latex):
    """将 LaTeX 公式通过 pandoc 转换为 OMML 并插入段落"""
    with tempfile.NamedTemporaryFile(mode='w', suffix='.tex', delete=False, encoding='utf-8') as f:
        f.write(f'\\documentclass{{article}}\\begin{{document}}${latex}$\\end{{document}}')
        tex_path = f.name
    docx_path = tex_path.replace('.tex', '.docx')
    subprocess.run(['pandoc', tex_path, '-o', docx_path], check=True, capture_output=True)
    tmp_doc = Document(docx_path)
    for p in tmp_doc.paragraphs:
        for child in p._element:
            if child.tag.endswith('}oMath') or child.tag.endswith('}oMathPara'):
                paragraph._element.append(child)
    os.unlink(tex_path)
    os.unlink(docx_path)


def insert_formula_centered(doc, latex):
    """插入居中公式"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    insert_formula(p, latex)
    return p


def init_styles(doc):
    """初始化 Heading 1/2 样式定义"""
    # Heading 1: 样式定义 sz=44，run 用 14pt 覆盖
    h1 = doc.styles['Heading 1']
    h1_rPr = h1.element.find(qn('w:rPr'))
    if h1_rPr is not None:
        for tag in ['w:sz', 'w:szCs']:
            el = h1_rPr.find(qn(tag))
            if el is not None:
                el.set(qn('w:val'), '44')
        color = h1_rPr.find(qn('w:color'))
        if color is not None:
            color.set(qn('w:val'), '000000')
            for a in ['w:themeColor', 'w:themeShade']:
                if color.get(qn(a)) is not None:
                    del color.attrib[qn(a)]
        for tag in ['w:b', 'w:bCs']:
            el = h1_rPr.find(qn(tag))
            if el is not None:
                h1_rPr.remove(el)
    h1_pPr = h1.element.find(qn('w:pPr'))
    if h1_pPr is not None:
        sp = h1_pPr.find(qn('w:spacing'))
        if sp is None:
            sp = OxmlElement('w:spacing')
            h1_pPr.append(sp)
        sp.set(qn('w:before'), str(H1_SPACING_BEFORE))
        sp.set(qn('w:after'), str(H1_SPACING_AFTER))
        sp.set(qn('w:line'), str(H1_SPACING_LINE))
        sp.set(qn('w:lineRule'), H1_SPACING_LINE_RULE)

    # Heading 2: 样式定义 sz=32，run 用 10.5pt 覆盖
    h2 = doc.styles['Heading 2']
    h2_rPr = h2.element.find(qn('w:rPr'))
    if h2_rPr is not None:
        for tag in ['w:sz', 'w:szCs']:
            el = h2_rPr.find(qn(tag))
            if el is not None:
                el.set(qn('w:val'), '32')
        color = h2_rPr.find(qn('w:color'))
        if color is not None:
            color.set(qn('w:val'), '000000')
            for a in ['w:themeColor', 'w:themeShade']:
                if color.get(qn(a)) is not None:
                    del color.attrib[qn(a)]
    h2_pPr = h2.element.find(qn('w:pPr'))
    if h2_pPr is not None:
        sp = h2_pPr.find(qn('w:spacing'))
        if sp is None:
            sp = OxmlElement('w:spacing')
            h2_pPr.append(sp)
        sp.set(qn('w:before'), str(H2_SPACING_BEFORE))
        sp.set(qn('w:after'), str(H2_SPACING_AFTER))
        sp.set(qn('w:line'), str(H2_SPACING_LINE))
        sp.set(qn('w:lineRule'), H2_SPACING_LINE_RULE)

    # Normal 样式：sz=21 (10.5pt), 两端对齐
    normal = doc.styles['Normal']
    normal_rPr = normal.element.find(qn('w:rPr'))
    if normal_rPr is not None:
        for tag in ['w:sz', 'w:szCs']:
            el = normal_rPr.find(qn(tag))
            if el is not None:
                el.set(qn('w:val'), '21')
    normal_pPr = normal.element.find(qn('w:pPr'))
    if normal_pPr is not None:
        jc = normal_pPr.find(qn('w:jc'))
        if jc is None:
            jc = OxmlElement('w:jc')
            normal_pPr.append(jc)
        jc.set(qn('w:val'), 'both')


def _get_image_rel_ids(element):
    """从 XML 元素中提取图片关系 ID"""
    rel_ids = []
    for blip in element.iter('{http://schemas.openxmlformats.org/drawingml/2006/main}blip'):
        rid = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
        if rid:
            rel_ids.append(rid)
    return rel_ids


def _replace_rel_id(element, old_rid, new_rid):
    """替换 XML 元素中的关系 ID"""
    ns_r = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
    for blip in element.iter('{http://schemas.openxmlformats.org/drawingml/2006/main}blip'):
        if blip.get(f'{{{ns_r}}}embed') == old_rid:
            blip.set(f'{{{ns_r}}}embed', new_rid)


# ============================================================
# 生成随机测试文章
# ============================================================
doc = Document()
init_styles(doc)

# 移除默认空段落
if doc.paragraphs:
    el = doc.paragraphs[0]._element
    el.getparent().remove(el)

# ===== Section 0: 封面（从封面.docx 原样复制）=====
setup_section(doc.sections[0], cols=1)

# 读取封面文档并复制内容（含图片）
cover_path = os.path.join(os.path.dirname(__file__), '封面.docx')
if os.path.exists(cover_path):
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    cover_doc = Document(cover_path)
    body = doc.element.body
    # 获取 body 的第一个子元素，用于插入到开头
    first_child = body[0] if len(body) > 0 else None
    # 从后往前遍历封面元素，保持顺序地插入到开头
    cover_elements = []
    for element in cover_doc.element.body:
        tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag
        if tag in ('p', 'tbl'):
            cover_elements.append(element)
    # 按顺序插入到 body 开头
    for element in cover_elements:
        tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag
        if tag == 'p':
            new_p = copy.deepcopy(element)
            # 复制图片关系
            for rel_id in _get_image_rel_ids(element):
                rel = cover_doc.part.rels[rel_id]
                if "image" in rel.reltype:
                    img_part = rel.target_part
                    new_rid = doc.part.relate_to(img_part, RT.IMAGE)
                    _replace_rel_id(new_p, rel_id, new_rid)
            if first_child is not None:
                first_child.addprevious(new_p)
            else:
                body.append(new_p)
        elif tag == 'tbl':
            new_tbl = copy.deepcopy(element)
            if first_child is not None:
                first_child.addprevious(new_tbl)
            else:
                body.append(new_tbl)

# 封面后分节
sec_cover_end = doc.add_section()
setup_section(sec_cover_end, cols=1)
set_continuous(sec_cover_end)

# ===== Section 1: 标题 + 作者 + 摘要（单栏）=====
# 标题（18pt，Heading 1 样式，居中）
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.style = doc.styles['Heading 1']
# 段落级别：before=0, after=0, line=0/atLeast（模板实测）
set_spacing(p, before=0, after=0, line=0, rule='atLeast')
run1 = p.add_run('基于深度强化学习的')
set_font(run1, '黑体', 'Times New Roman', SIZE_TITLE)
run2 = p.add_run('多机器人协同路径规划方法研究')
set_font(run2, '黑体', 'Times New Roman', SIZE_TITLE)

# 作者（14pt，居中，before=157）
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_spacing(p, before=AUTHOR_SPACING_BEFORE, line=0, rule='atLeast')
run = p.add_run('张三    李四    王五')
set_font(run, '宋体', 'Times New Roman', SIZE_AUTHOR, bold=True)

# 单位（9pt，居中，after=157）
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_spacing(p, after=AFFILIATION_SPACING_AFTER, line=0, rule='atLeast')
run = p.add_run('（郑州大学 电气与信息工程学院，河南 郑州 450001）')
set_font(run, '宋体', 'Times New Roman', SIZE_AFFILIATION)

# 中文摘要（9pt，left=351twips, right=390twips）
p = doc.add_paragraph()
set_indent(p, left=IND_ABSTRACT_LEFT, right=IND_ABSTRACT_RIGHT)
label = p.add_run('摘  要：')
set_font(label, '黑体', 'Times New Roman', SIZE_ABSTRACT, bold=True)
content = p.add_run(
    '针对多机器人系统在复杂环境中的协同路径规划问题，本文提出了一种基于深度强化学习的分布式协同规划方法。'
    '该方法通过设计一种改进的多智能体近端策略优化算法，使各机器人能够在不依赖全局信息的条件下，'
    '仅通过局部观测和有限通信实现高效的协同路径规划。实验结果表明，所提方法在100×100的栅格环境中，'
    '相比传统A*算法和遗传算法，路径长度缩短了15.3%，规划时间减少了42.7%，'
    '且在动态障碍物场景下表现出更强的鲁棒性。'
)
set_font(content, '宋体', 'Times New Roman', SIZE_ABSTRACT)

# 关键词
p = doc.add_paragraph()
set_indent(p, left=IND_ABSTRACT_LEFT, right=IND_ABSTRACT_RIGHT)
label = p.add_run('关键词：')
set_font(label, '黑体', 'Times New Roman', SIZE_ABSTRACT, bold=True)
content = p.add_run('深度强化学习；多机器人系统；路径规划；多智能体协同；近端策略优化')
set_font(content, '宋体', 'Times New Roman', SIZE_ABSTRACT)

# ===== Section 2: 双栏正文 =====
sec2 = doc.add_section()
setup_section(sec2, cols=2)
set_continuous(sec2)

# 0 引言
make_heading_1(doc, '0 引言')

make_body(doc,
    '随着机器人技术的快速发展，多机器人系统在物流配送、搜索救援、环境监测等领域的应用日益广泛。'
    '路径规划作为多机器人系统的核心问题之一，直接影响着系统的整体效率和安全性。'
    '传统的集中式规划方法虽然能够获得全局最优解，但随着机器人数量的增加，计算复杂度呈指数增长，'
    '难以满足实时性要求。')

make_body(doc,
    '近年来，深度强化学习在单机器人路径规划中取得了显著成果，但在多机器人协同场景中仍面临诸多挑战。'
    '首先，多机器人系统的状态空间和动作空间维度随机器人数量急剧增长；'
    '其次，各机器人之间的通信和协调机制设计直接影响协同效果；'
    '最后，如何在保证规划质量的同时提高计算效率是实际部署的关键问题。')

make_body(doc,
    '本文针对上述挑战，提出了一种基于改进多智能体近端策略优化（MAPPO）算法的分布式协同路径规划方法。'
    '该方法的核心创新包括：设计了一种基于注意力机制的局部信息融合模块，'
    '使各机器人能够有效利用邻居信息；引入了课程学习策略，逐步增加训练难度以提高学习效率；'
    '提出了一种动态奖励函数，平衡路径长度、安全性和协同效率。')

# 1 相关工作
make_heading_1(doc, '1 相关工作')

make_heading_2(doc, '1.1 传统路径规划方法')

make_body(doc,
    '传统路径规划方法主要分为基于搜索的方法和基于采样的方法。'
    'A*算法作为经典的基于搜索的方法，通过启发式函数引导搜索方向，能够在静态环境中找到最优路径。'
    '然而，在多机器人场景中，A*算法需要为每个机器人单独规划路径，难以处理机器人之间的冲突。'
    'RRT（快速扩展随机树）算法作为基于采样的方法，适用于高维空间的路径规划，'
    '但其随机性导致路径质量不稳定，且难以保证最优性。')

make_heading_2(doc, '1.2 基于强化学习的路径规划')

make_body(doc,
    '深度强化学习为路径规划提供了新的思路。DQN（深度Q网络）算法通过将状态映射到Q值函数，'
    '能够在高维状态空间中学习有效的策略。PPO（近端策略优化）算法通过限制策略更新幅度，'
    '提高了训练的稳定性。在多智能体领域，MAPPO算法将PPO扩展到多智能体场景，'
    '通过共享价值网络和独立策略网络实现协同学习。')

make_body(doc,
    '然而，现有的基于强化学习的多机器人路径规划方法存在以下不足：'
    '大多数方法假设机器人之间可以进行完全通信，这在实际场景中往往不现实；'
    '奖励函数的设计通常只考虑路径长度，忽略了安全性和协同效率；'
    '训练过程缺乏有效的课程学习策略，导致学习效率低下。')

# 2 方法
make_heading_1(doc, '2 方法')

make_heading_2(doc, '2.1 问题建模')

make_body(doc,
    '本文将多机器人协同路径规划问题建模为一个去中心化的部分可观测马尔可夫决策过程。'
    '每个机器人作为一个智能体，其观测范围限定在以自身为中心的局部区域内。'
    '状态空间包括机器人自身位置、目标位置、局部障碍物信息和邻居机器人信息。'
    '动作空间为离散的八方向移动加上原地不动。')

# MDP 公式
insert_formula_centered(doc, r'\langle \mathcal{S}, \mathcal{A}, \mathcal{O}, \mathcal{T}, \mathcal{R}, \gamma \rangle')

make_heading_2(doc, '2.2 改进的MAPPO算法')

make_body(doc,
    '本文对标准MAPPO算法进行了三方面改进。第一，设计了基于多头注意力机制的局部信息融合模块，'
    '使每个机器人能够动态权衡不同邻居信息的重要性。第二，引入了课程学习策略，'
    '从简单的2机器人场景逐步扩展到10机器人场景，提高了训练效率。'
    '第三，提出了综合考虑路径长度、碰撞风险和协同效率的动态奖励函数。')

make_body(doc,
    '具体而言，局部信息融合模块的计算过程如下：首先将每个邻居的观测信息编码为特征向量，'
    '然后通过多头注意力机制计算当前机器人对各邻居的关注权重，'
    '最后将加权融合后的特征与自身特征拼接作为策略网络的输入。'
    '这种设计使得机器人能够根据任务需求动态调整对不同邻居的关注程度。')

# 动态奖励函数公式
insert_formula_centered(doc, r'R_t = \alpha \cdot R_t^{\text{dist}} + \beta \cdot R_t^{\text{safe}} + \gamma \cdot R_t^{\text{coop}}')

make_body(doc,
    '其中，动态奖励函数如公式所示，分别对路径距离、安全性和协同效率进行加权求和，'
    '权重系数\\alpha、\\beta和\\gamma根据训练阶段动态调整。')

# 3 实验
make_heading_1(doc, '3 实验')

make_heading_2(doc, '3.1 实验设置')

make_body(doc,
    '实验在100×100的栅格环境中进行，机器人数量从2到10不等。'
    '障碍物密度设置为10%、20%和30%三个等级。'
    '每个机器人需要从随机起点移动到随机终点，同时避免与障碍物和其他机器人发生碰撞。'
    '训练过程使用Adam优化器，学习率为3×10⁻⁴，折扣因子为0.99。'
    '每个场景训练100万步，每1000步进行一次评估。')

make_heading_2(doc, '3.2 实验结果')

make_body(doc,
    '实验结果表明，本文方法在各项指标上均优于基线方法。'
    '在路径长度方面，相比A*算法缩短了15.3%，相比遗传算法缩短了22.1%。'
    '在规划时间方面，相比A*算法减少了42.7%，相比遗传算法减少了58.3%。'
    '在成功率方面，本文方法在所有场景中均保持95%以上的成功率，'
    '而A*算法在高障碍物密度场景中成功率下降到78%。')

make_body(doc,
    '消融实验验证了各改进模块的有效性。去除注意力机制后，路径长度增加8.2%；'
    '去除课程学习后，训练时间增加35.6%；使用静态奖励函数后，成功率下降12.4%。'
    '这表明三个改进模块都对最终性能有显著贡献。')

# PPO 损失函数公式
make_body(doc, '本文采用的PPO损失函数定义为：')
insert_formula_centered(doc, r'L^{\text{CLIP}}(\theta) = \mathbb{E}_t \left[ \min \left( r_t(\theta) \hat{A}_t, \ \text{clip}(r_t(\theta), 1-\epsilon, 1+\epsilon) \hat{A}_t \right) \right]')

# 4 结论
make_heading_1(doc, '4 结论')

make_body(doc,
    '本文提出了一种基于深度强化学习的多机器人协同路径规划方法。'
    '通过设计改进的MAPPO算法、注意力机制融合模块和动态奖励函数，'
    '实现了在复杂环境中的高效协同规划。实验结果验证了所提方法的有效性和优越性。'
    '未来工作将探索更复杂的动态环境和异构机器人系统的协同规划问题。')

# ===== Section 3: 参考文献（双栏）=====
sec3 = doc.add_section()
setup_section(sec3, cols=2)
set_continuous(sec3)

make_heading_1(doc, '参考文献')

# 中文参考文献
make_ref_zh(doc, '[1] 王雪芳,郑建岚,罗素蓉,等.矿掺合料对混凝土电阻率的影响[J].福州大学学报（自然科学版）, 2008, 36(3):416-421.')
make_ref_zh(doc, '[2] 谢友均,马昆林,龙广成,等.矿物掺合料对混凝土中氯离子渗透性的影响[J]. 硅酸盐学报,2006, 34(11):1395-1400.')
make_ref_zh(doc, '[3] 赵卓,蒋晓东.受腐蚀混凝土结构耐久性检测诊断[M].郑州: 黄河水利出版社, 2006.')

# 英文参考文献
make_ref_en(doc, '[4] RAIF B A, BEKIR T L. Influence of fly ash on corrosion resistance of concrete[J]. Construction and Building Materials, 2013, 42: 1-8.')
make_ref_en(doc, '[5] XIE Y J, MA K L, LONG G C, et al. Effect of mineral admixtures on chloride ion permeability of concrete[J]. Journal of the Chinese Ceramic Society, 2006, 34(11):1395-1400.')
make_ref_en(doc, '[6] DONG X Q, BAI X H, WU Z A, et al. Study on the application of resistivity technology in concrete steel corrosion testing[J]. Concrete, 2007(9):9-12.')
make_ref_en(doc, '[7] SEMION Z, Konstantin K. Effect of internal curing on durability of concrete[J]. Cement and Concrete Research, 2015, 78: 1-12.')
make_ref_en(doc, '[8] GIROU E M. Computer algebra in scientific computing[C]//Proceedings of CASC, 2010: 1-12.')

# ===== Section 4: 英文摘要（单栏）=====
sec4 = doc.add_section()
setup_section(sec4, cols=1)

# 英文标题（12pt，Heading 1 样式，left align，before=0 after=0）
p = doc.add_paragraph()
p.style = doc.styles['Heading 1']
set_spacing(p, before=0, after=0, line=240, rule='auto')
run = p.add_run('Deep Reinforcement Learning Based Multi-Robot Cooperative Path Planning')
set_font(run, 'Times New Roman', 'Times New Roman', SIZE_ENG_TITLE)

# 英文作者（11pt，居中）
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Zhang San    Li Si    Wang Wu')
set_font(run, 'Times New Roman', 'Times New Roman', SIZE_ENG_AUTHOR)

# 英文单位（9pt，居中）
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('(School of Electrical and Information Engineering, Zhengzhou University, Zhengzhou 450001, China)')
set_font(run, 'Times New Roman', 'Times New Roman', SIZE_AFFILIATION)

# 英文摘要（9pt，无缩进）
p = doc.add_paragraph()
label = p.add_run('Abstract: ')
set_font(label, 'Times New Roman', 'Times New Roman', SIZE_ABSTRACT, bold=True)
content = p.add_run(
    'Aiming at the cooperative path planning problem of multi-robot systems in complex environments, '
    'this paper proposes a distributed cooperative planning method based on deep reinforcement learning. '
    'The method designs an improved multi-agent proximal policy optimization algorithm that enables each robot '
    'to achieve efficient cooperative path planning through local observations and limited communication '
    'without relying on global information. Experimental results show that in a 100×100 grid environment, '
    'the proposed method reduces path length by 15.3% and planning time by 42.7% compared to A* algorithm '
    'and genetic algorithm, and exhibits stronger robustness in dynamic obstacle scenarios.'
)
set_font(content, 'Times New Roman', 'Times New Roman', SIZE_ABSTRACT)

# 英文关键词（9pt）
p = doc.add_paragraph()
label = p.add_run('Keywords: ')
set_font(label, 'Times New Roman', 'Times New Roman', SIZE_ABSTRACT, bold=True)
content = p.add_run(
    'deep reinforcement learning; multi-robot system; path planning; multi-agent cooperation; proximal policy optimization'
)
set_font(content, 'Times New Roman', 'Times New Roman', SIZE_ABSTRACT)

# ===== 保存 =====
output_path = r'E:\Code\Embed\作业\智能机器人控制\测试文章-郑州大学工学报格式.docx'
doc.save(output_path)
print(f'已保存: {output_path}')
