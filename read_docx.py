from docx import Document
doc = Document(r'E:\Code\Embed\作业\智能机器人控制\马明亮-202522852018965.docx')

# Print paragraph styles and text
for i, para in enumerate(doc.paragraphs[:100]):
    style = para.style.name if para.style else 'None'
    text = para.text[:200] if para.text else ''
    if text.strip():
        print(f'[{i}] Style:{style} | {text}')

print('\n--- TABLES ---')
for ti, table in enumerate(doc.tables[:5]):
    print(f'Table {ti}: {len(table.rows)} rows x {len(table.columns)} cols')
    for ri, row in enumerate(table.rows[:3]):
        cells = [c.text[:50] for c in row.cells]
        print(f'  Row {ri}: {cells}')
