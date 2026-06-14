"""
郑州大学工学报格式转换脚本模板 v1.1
将源Word文档重新格式化为郑州大学工学报格式

使用方法：
1. 修改 SOURCE_PATH 为源Word文档路径
2. 修改 OUTPUT_PATH 为输出文档路径
3. 根据源文档结构调整 read_source_doc() 中的识别逻辑
4. 运行: E:/Anaconda/python.exe generate_template.py
"""

import os
import re
from docx import Document
from docx.shared import Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============================================================
# 配置区域
# ============================================================
SOURCE_PATH = r'源文档.docx'
OUTPUT_PATH = r'输出文档.docx'
FIG_DIR = r'figures'  # 图片目录（可选）

# ============================================================
# 格式常量（郑州大学工学报）
# ============================================================
SIZE_AUTHOR = 177800        # 14pt (sz=28)
SIZE_AFFILIATION = 114300   # 9pt (sz=18)
SIZE_HEADING1 = 177800      # 14pt (sz=28)
SIZE_HEADING2 = 133350      # 11pt (sz=21)
SIZE_BODY = 133350          # 10.5pt (sz=21)
SIZE_ABSTRACT = 114300      # 9pt (sz=18)
SIZE_CAPTION = 114300       # 9pt (sz=18)
SIZE_REF = 133350           # 10.5pt (sz=21)
SIZE_ENG_TITLE = 152400     # 12pt (sz=24)

FI_BODY = 266700
RI_BODY = 255905

PAGE_W = 11906
PAGE_H = 16838
MARGIN_TB = 1440
MARGIN_LR = 1800
COL_SPACE = 425


# ============================================================
# 工具函数
# ============================================================
def set_font(run, ea='宋体', fn='Times New Roman', size=None, bold=False):
    run.font.name = fn
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    if size:
        run.font.size = Emu(size)
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), ea)
    rFonts.set(qn('w:ascii'), fn)
    rFonts.set(qn('w:hAnsi'), fn)


def set_spacing(para, line=None, rule=None):
    pPr = para._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        para._element.insert(0, pPr)
    sp = pPr.find(qn('w:spacing'))
    if sp is None:
        sp = OxmlElement('w:spacing')
        pPr.append(sp)
    if line is not None:
        sp.set(qn('w:line'), str(line))
    if rule is not None:
        sp.set(qn('w:lineRule'), rule)


def set_indent(para, first=None, left=None, right=None):
    pPr = para._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        para._element.insert(0, pPr)
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind')
        pPr.append(ind)
    if first is not None:
        ind.set(qn('w:firstLine'), str(int(first / 635)))
    if left is not None:
        ind.set(qn('w:left'), str(int(left / 635)))
    if right is not None:
        ind.set(qn('w:right'), str(int(right / 635)))


def set_continuous(section):
    sectPr = section._sectPr
    type_el = sectPr.find(qn('w:type'))
    if type_el is None:
        type_el = OxmlElement('w:type')
        pgSz = sectPr.find(qn('w:pgSz'))
        if pgSz is not None:
            pgSz.addprevious(type_el)
        else:
            sectPr.append(type_el)
    type_el.set(qn('w:val'), 'continuous')


def setup_section(section, cols=1, col_space=COL_SPACE):
    sectPr = section._sectPr
    pgSz = sectPr.find(qn('w:pgSz'))
    if pgSz is None:
        pgSz = OxmlElement('w:pgSz')
        sectPr.append(pgSz)
    pgSz.set(qn('w:w'), str(PAGE_W))
    pgSz.set(qn('w:h'), str(PAGE_H))
    pgMar = sectPr.find(qn('w:pgMar'))
    if pgMar is None:
        pgMar = OxmlElement('w:pgMar')
        sectPr.append(pgMar)
    pgMar.set(qn('w:top'), str(MARGIN_TB))
    pgMar.set(qn('w:bottom'), str(MARGIN_TB))
    pgMar.set(qn('w:left'), str(MARGIN_LR))
    pgMar.set(qn('w:right'), str(MARGIN_LR))
    old_cols = sectPr.findall(qn('w:cols'))
    for c in old_cols:
        sectPr.remove(c)
    cols_elem = OxmlElement('w:cols')
    if cols > 1:
        cols_elem.set(qn('w:num'), str(cols))
        cols_elem.set(qn('w:space'), str(col_space))
    sectPr.append(cols_elem)


def make_heading_1(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles['Heading 1']
    m = re.match(r'^(\d+\s+)(.+)$', text)
    if m:
        run1 = p.add_run(m.group(1))
        set_font(run1, '宋体', 'Times New Roman', SIZE_HEADING1)
        run2 = p.add_run(m.group(2))
        set_font(run2, '黑体', 'Times New Roman', SIZE_HEADING1)
    else:
        run = p.add_run(text)
        set_font(run, '黑体', 'Times New Roman', SIZE_HEADING1)
    set_spacing(p, 240, 'auto')
    return p


def make_heading_2(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles['Heading 2']
    run = p.add_run(text)
    set_font(run, '黑体', 'Times New Roman', SIZE_HEADING2)
    set_spacing(p, 240, 'auto')
    return p


def make_body(doc, text):
    p = doc.add_paragraph()
    set_indent(p, first=FI_BODY, right=RI_BODY)
    run = p.add_run(text)
    set_font(run, '宋体', 'Times New Roman', SIZE_BODY)
    return p


def make_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, '黑体', 'Times New Roman', SIZE_CAPTION, bold=True)
    return p


def make_image(doc, path, width_in=3.2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Cm(width_in * 2.54))
    return p


def make_ref(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, '宋体', 'Times New Roman', SIZE_REF)
    return p


# ============================================================
# 样式初始化
# ============================================================
def init_styles(doc):
    # Heading 1: sz=44 (22pt), 黑色, 不加粗, 段前段后0
    h1 = doc.styles['Heading 1']
    h1_rPr = h1.element.find(qn('w:rPr'))
    if h1_rPr is not None:
        sz = h1_rPr.find(qn('w:sz'))
        if sz is not None:
            sz.set(qn('w:val'), '44')
        szCs = h1_rPr.find(qn('w:szCs'))
        if szCs is not None:
            szCs.set(qn('w:val'), '44')
        color = h1_rPr.find(qn('w:color'))
        if color is not None:
            color.set(qn('w:val'), '000000')
            for a in ['w:themeColor', 'w:themeShade']:
                if color.get(qn(a)) is not None:
                    del color.attrib[qn(a)]
        for tag in ['w:b', 'w:bCs']:
            el = h1_rPr.find(qn(tag))
            if el is not None:
                h1_rPr.remove(el)
    # Heading 1 段间距: 段前段后0
    h1_pPr = h1.element.find(qn('w:pPr'))
    if h1_pPr is not None:
        sp = h1_pPr.find(qn('w:spacing'))
        if sp is None:
            sp = OxmlElement('w:spacing')
            h1_pPr.append(sp)
        sp.set(qn('w:before'), '0')
        sp.set(qn('w:after'), '0')

    # Heading 2: 黑色, 段前段后0
    h2 = doc.styles['Heading 2']
    h2_rPr = h2.element.find(qn('w:rPr'))
    if h2_rPr is not None:
        color = h2_rPr.find(qn('w:color'))
        if color is not None:
            color.set(qn('w:val'), '000000')
            for a in ['w:themeColor', 'w:themeShade']:
                if color.get(qn(a)) is not None:
                    del color.attrib[qn(a)]
    h2_pPr = h2.element.find(qn('w:pPr'))
    if h2_pPr is not None:
        sp = h2_pPr.find(qn('w:spacing'))
        if sp is not None:
            sp.set(qn('w:before'), '0')
            sp.set(qn('w:after'), '0')


# ============================================================
# 封面页拼接
# ============================================================
def find_cover_docx(work_dir):
    """在工作目录下查找封面docx文件（文件名含'封面'）"""
    for f in os.listdir(work_dir):
        if f.endswith('.docx') and '封面' in f and not f.startswith('~'):
            return os.path.join(work_dir, f)
    return None


def merge_cover_page(doc, cover_path):
    """将封面docx的所有内容拼接到文档开头，保持原始格式，正确处理图片"""
    cover = Document(cover_path)
    body = doc.element.body

    # 提取封面的图片数据，建立 old_rId -> (image_bytes, content_type) 映射
    cover_images = {}
    for rel_id, rel in cover.part.rels.items():
        if 'image' in rel.reltype and hasattr(rel.target_part, 'blob'):
            cover_images[rel_id] = (rel.target_part.blob, rel.target_part.content_type)

    # 获取文档当前第一个元素
    first_element = body[0]

    # 将封面的每个元素复制到新文档的最前面
    for element in cover.element.body:
        if element.tag.endswith('}sectPr'):
            continue
        first_element.addprevious(element)

    # 修复图片引用：找到封面段落中的 drawing blip，重新注册图片到新文档
    import io
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from docx.opc.part import Part as OpcPart
    from docx.opc.packuri import PackURI

    for elem in body:
        if elem == first_element:
            break
        if elem.tag.endswith('}p'):
            blips = elem.findall('.//' + qn('a:blip'))
            for blip in blips:
                old_embed = blip.get(qn('r:embed'))
                if old_embed and old_embed in cover_images:
                    img_bytes, img_ct = cover_images[old_embed]
                    ext_map = {
                        'image/png': 'png', 'image/jpeg': 'jpg', 'image/jpg': 'jpg',
                        'image/gif': 'gif', 'image/bmp': 'bmp',
                        'image/x-emf': 'emf', 'image/x-wmf': 'wmf',
                    }
                    ext = ext_map.get(img_ct, 'png')
                    # WMF → PNG
                    if 'wmf' in img_ct:
                        try:
                            from PIL import Image
                            img = Image.open(io.BytesIO(img_bytes))
                            png_buf = io.BytesIO()
                            img.save(png_buf, 'PNG')
                            img_bytes = png_buf.getvalue()
                            img_ct = 'image/png'
                            ext = 'png'
                        except Exception:
                            pass
                    # 计算新 rId
                    max_num = 0
                    for rid in doc.part.rels.keys():
                        if rid.startswith('rId'):
                            try:
                                n = int(rid[3:])
                                max_num = max(max_num, n)
                            except ValueError:
                                pass
                    new_rid = f'rId{max_num + 1}'
                    img_partname = PackURI(f'/word/media/cover_{max_num + 1}.{ext}')
                    img_part = OpcPart(img_partname, img_ct, img_bytes, doc.part.package)
                    doc.part.relate_to(img_part, RT.IMAGE)
                    blip.set(qn('r:embed'), new_rid)

    # 在封面内容之后插入分节符
    new_sect = OxmlElement('w:sectPr')
    pgSz = OxmlElement('w:pgSz')
    pgSz.set(qn('w:w'), str(PAGE_W))
    pgSz.set(qn('w:h'), str(PAGE_H))
    new_sect.append(pgSz)
    first_element.addprevious(new_sect)


# ============================================================
# 读取源文档
# ============================================================
def read_source_doc(path):
    """读取源Word文档，提取结构化内容

    注意：需要根据实际源文档的结构调整识别逻辑。
    以下是通用的基于样式名的识别方式。
    """
    src = Document(path)
    content = {
        'title': '',
        'authors': '',
        'affiliation': '',
        'abstract_zh': '',
        'keywords_zh': '',
        'abstract_en': '',
        'keywords_en': '',
        'sections': [],      # [(type, text), ...] type: h1/h2/body/caption
        'references': [],
    }

    for para in src.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name if para.style else 'None'

        # 根据样式名判断类型
        if style == 'Heading 1':
            content['sections'].append(('h1', text))
        elif style == 'Heading 2':
            content['sections'].append(('h2', text))
        elif '摘要' in text and len(text) < 10:
            continue  # 跳过摘要标签（会在生成时重建）
        elif '关键词' in text and len(text) < 10:
            continue
        else:
            content['sections'].append(('body', text))

    return content


# ============================================================
# 标题段落（段前段后0）
# ============================================================
def make_title(doc, title):
    """论文标题：Heading 1 样式，22pt 黑体，居中，段前段后0"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.style = doc.styles['Heading 1']
    # 段落级别显式设置段前段后为0
    pPr = p._element.find(qn('w:pPr'))
    sp = pPr.find(qn('w:spacing'))
    if sp is None:
        sp = OxmlElement('w:spacing')
        pPr.append(sp)
    sp.set(qn('w:before'), '0')
    sp.set(qn('w:after'), '0')
    run = p.add_run(title)
    set_font(run, '黑体', 'Times New Roman')  # 继承样式 sz=44 (22pt)
    return p


# ============================================================
# 主流程
# ============================================================
def main():
    # 0. 用户信息（使用 AskUserQuestion 采集）
    AUTHOR_NAME = '用户姓名'           # ← 替换为采集到的姓名
    STUDENT_ID = '用户学号'            # ← 替换为采集到的学号
    CLASS_INFO = '年级+专业'           # ← 替换为采集到的班级
    AFFILIATION = '（郑州大学 电气与信息工程学院，河南 郑州 450001）'

    # 1. 读取源文档，自动抓取标题
    source = read_source_doc(SOURCE_PATH)
    title_zh = source.get('title', '')  # 自动从源文档提取中文标题
    title_en = source.get('title_en', '')  # 自动提取英文标题

    # 2. 创建新文档
    doc = Document()
    init_styles(doc)
    if doc.paragraphs:
        el = doc.paragraphs[0]._element
        el.getparent().remove(el)

    # 3. 封面页拼接（可选：同目录下有含"封面"的docx则自动拼接）
    work_dir = os.path.dirname(os.path.abspath(SOURCE_PATH))
    cover_path = find_cover_docx(work_dir)
    if cover_path:
        print(f'Found cover: {cover_path}')
        merge_cover_page(doc, cover_path)
    else:
        print('No cover file found, skipping.')

    # 4. Section 0: 标题 + 作者 + 摘要（单栏）
    setup_section(doc.sections[0], cols=1)

    # 标题（自动抓取，段前段后0）
    make_title(doc, title_zh)

    # 作者（用户提供姓名）
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(AUTHOR_NAME)
    set_font(run, '宋体', 'Times New Roman', SIZE_AUTHOR, bold=True)

    # 单位
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(AFFILIATION)
    set_font(run, '宋体', 'Times New Roman', SIZE_AFFILIATION)

    # 中文摘要
    p = doc.add_paragraph()
    set_indent(p, left=230505, right=255905)
    label = p.add_run('摘  要：')
    set_font(label, '黑体', 'Times New Roman', SIZE_ABSTRACT, bold=True)
    content = p.add_run(source.get('abstract_zh', '在此输入中文摘要...'))
    set_font(content, '宋体', 'Times New Roman', SIZE_ABSTRACT)

    # 关键词
    p = doc.add_paragraph()
    set_indent(p, left=230505, right=255905)
    label = p.add_run('关键词：')
    set_font(label, '黑体', 'Times New Roman', SIZE_ABSTRACT, bold=True)
    content = p.add_run(source.get('keywords_zh', '关键词1；关键词2'))
    set_font(content, '宋体', 'Times New Roman', SIZE_ABSTRACT)

    # 5. Section 1: 正文（双栏，continuous）
    sec1 = doc.add_section()
    setup_section(sec1, cols=2)
    set_continuous(sec1)

    # 按源文档结构写入正文
    for stype, text in source['sections']:
        if stype == 'h1':
            make_heading_1(doc, text)
        elif stype == 'h2':
            make_heading_2(doc, text)
        elif stype == 'caption':
            make_caption(doc, text)
        else:
            make_body(doc, text)

    # 6. Section 2: 参考文献（双栏，continuous）
    sec2 = doc.add_section()
    setup_section(sec2, cols=2)
    set_continuous(sec2)
    make_heading_1(doc, '参考文献')
    for ref in source.get('references', []):
        make_ref(doc, ref)

    # 7. Section 3: 英文摘要（单栏，新页）
    sec3 = doc.add_section()
    setup_section(sec3, cols=1)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('English Title')
    set_font(run, 'Times New Roman', 'Times New Roman', SIZE_ENG_TITLE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Author Name')
    set_font(run, 'Times New Roman', 'Times New Roman', SIZE_AUTHOR, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('(University, City, Country)')
    set_font(run, 'Times New Roman', 'Times New Roman', SIZE_AFFILIATION)

    p = doc.add_paragraph()
    set_indent(p, right=RI_BODY)
    label = p.add_run('Abstract: ')
    set_font(label, 'Times New Roman', 'Times New Roman', SIZE_ABSTRACT, bold=True)
    content = p.add_run(source.get('abstract_en', 'Enter English abstract here...'))
    set_font(content, 'Times New Roman', 'Times New Roman', SIZE_ABSTRACT)

    p = doc.add_paragraph()
    label = p.add_run('Keywords: ')
    set_font(label, 'Times New Roman', 'Times New Roman', SIZE_ABSTRACT, bold=True)
    content = p.add_run(source.get('keywords_en', 'keyword1; keyword2'))
    set_font(content, 'Times New Roman', 'Times New Roman', SIZE_ABSTRACT)

    # 8. 保存
    doc.save(OUTPUT_PATH)
    print(f'Saved: {OUTPUT_PATH}')


if __name__ == '__main__':
    main()
