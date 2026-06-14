# -*- coding: utf-8 -*-
"""
Generate TD-MPC2 Chinese translation in 郑州大学工学报 format.
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_font(run, name_cn='宋体', name_en='Times New Roman', size=None, bold=False):
    """Set font for a run."""
    run.font.name = name_en
    run.font.bold = bold
    if size:
        run.font.size = Pt(size)
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = r.makeelement(qn('w:rPr'), {})
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), name_cn)

def add_heading_1(doc, text):
    """Add Heading 1 style paragraph."""
    p = doc.add_paragraph()
    p.style = doc.styles['Heading 1']
    run = p.add_run(text)
    set_font(run, '黑体', 'Times New Roman', 14, bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def add_heading_2(doc, text):
    """Add Heading 2 style paragraph."""
    p = doc.add_paragraph()
    p.style = doc.styles['Heading 2']
    run = p.add_run(text)
    set_font(run, '黑体', 'Times New Roman', 12, bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def add_normal(doc, text, first_line_indent=True, font_size=10.5):
    """Add normal paragraph with proper formatting."""
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run(text)
    set_font(run, '宋体', 'Times New Roman', font_size)
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    return p

def add_title(doc, text):
    """Add paper title."""
    p = doc.add_paragraph()
    p.style = doc.styles['Heading 1']
    run = p.add_run(text)
    set_font(run, '黑体', 'Times New Roman', 16, bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p

def add_author(doc, text):
    """Add author name."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, '宋体', 'Times New Roman', 12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p

def add_affiliation(doc, text):
    """Add affiliation."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, '宋体', 'Times New Roman', 10.5)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p

def add_abstract(doc, text):
    """Add abstract."""
    p = doc.add_paragraph()
    label = p.add_run('摘  要：')
    set_font(label, '黑体', 'Times New Roman', 10.5, bold=True)
    content = p.add_run(text)
    set_font(content, '宋体', 'Times New Roman', 10.5)
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    return p

def add_keywords(doc, text):
    """Add keywords."""
    p = doc.add_paragraph()
    label = p.add_run('关键词：')
    set_font(label, '黑体', 'Times New Roman', 10.5, bold=True)
    content = p.add_run(text)
    set_font(content, '宋体', 'Times New Roman', 10.5)
    p.paragraph_format.line_spacing = 1.5
    return p


doc = Document()

# ===== 论文标题 =====
add_title(doc, 'TD-MPC2：面向连续控制的可扩展鲁棒世界模型')

# ===== 作者 =====
add_author(doc, 'Nicklas Hansen    Hao Su    Xiaolong Wang')

# ===== 单位 =====
add_affiliation(doc, '（加州大学圣地亚哥分校，美国 加利福尼亚州 圣地亚哥 92093）')

# ===== 摘要 =====
add_abstract(doc,
    'TD-MPC是一种基于模型的强化学习算法，它在学习到的隐式（无解码器）世界模型的潜在空间中进行局部轨迹优化。'
    '本文提出TD-MPC2，即对TD-MPC算法的一系列改进。实验表明，TD-MPC2在涵盖4个不同任务领域的104个在线强化学习任务上均明显优于基线方法，'
    '并且只需一组超参数即可取得稳定的表现。研究还发现，智能体的能力会随着模型和数据规模的增大而提升，'
    '作者成功训练了一个拥有3.17亿参数的单一智能体，使其能够在多个任务领域、多种机器人构型和动作空间中完成80项任务。'
    '最后，文章讨论了大规模TD-MPC2智能体在实际应用中的经验、机遇和潜在风险。'
)

# ===== 关键词 =====
add_keywords(doc, '基于模型的强化学习；世界模型；模型预测控制；连续控制；多任务学习')

# ===== 正文 =====

# 0 引言
add_heading_1(doc, '0 引言')

add_normal(doc,
    '在互联网规模的数据集上训练大型模型，已经在语言和视觉任务中催生了具有广泛能力的通用模型。'
    '这些模型的成功很大程度上得益于海量数据的可用性，以及能够随模型和数据规模可靠扩展的精心设计的架构。'
    '尽管近年来研究者已将这一范式推广到机器人领域，但要实现一个能够通过底层动作在多种机器人构型上学习执行多样化控制任务的通用具身智能体，'
    '并且所需数据来自大规模未经筛选（即质量参差不齐）的数据集，目前仍然是一个尚未解决的难题。'
)

add_normal(doc,
    '作者认为，当前通用具身智能体的研究路径存在两方面不足：其一，行为克隆方法通常假设训练数据为近似专家轨迹，'
    '这极大地限制了可用数据的规模；其二，缺少能够处理大规模未经筛选数据集的可扩展连续控制算法。'
)

add_normal(doc,
    '强化学习是从未经筛选的数据集中提取专家行为的理想框架。然而，现有的大多数强化学习算法都是为单任务学习设计的，'
    '依赖于针对每个任务单独调节的超参数，且缺乏选择这些超参数的系统性方法。'
    '一个能够处理大规模多任务数据集的算法，必然需要对不同任务之间的差异（如动作空间维度、探索难度和奖励分布）具有鲁棒性。'
    '本文提出的TD-MPC2正是朝着这一目标迈出的重要一步。'
    'TD-MPC2是一种基于模型的强化学习算法，专门设计用于在由多个任务领域、机器人构型和动作空间组成的大规模未经筛选数据集上学习通用世界模型，'
    '数据来源于覆盖不同技能水平的行为策略，且无需调节超参数。'
)

add_normal(doc,
    '该算法在TD-MPC的基础上发展而来，在学习到的隐式世界模型的潜在空间中进行局部轨迹优化。'
    '虽然TD-MPC系列算法在已有工作中展示了较强的实验性能，但大多数成果仅限于单任务学习，对可扩展性的关注较少。'
    '研究表明，简单地增大TD-MPC的模型和数据规模往往会导致智能体性能下降，这在强化学习文献中也较为常见。'
    '相比之下，TD-MPC2的扩展则能持续提升性能。'
    '作者的算法贡献主要体现在两个方面：通过重新审视核心设计选择来提高算法的鲁棒性，以及精心设计一种能够在不依赖领域知识的情况下适应多种机器人构型和动作空间数据集的架构。'
)

add_normal(doc,
    '作者在总共104个不同的连续控制任务上对TD-MPC2进行了评估，涵盖4个任务领域：DMControl、Meta-World、ManiSkill2和MyoSuite。'
    '这些任务涉及高维状态和动作空间（最高达39维动作空间）、图像观测、稀疏奖励、多物体操作、'
    '生理精确的肌肉骨骼运动控制、复杂运动（如四足和双足机器人构型），并且覆盖了广泛的任务难度范围。'
    '结果表明，TD-MPC2在所有任务中使用相同超参数的条件下，持续优于现有的基于模型和无模型方法。'
    '研究还发现，智能体的能力会随着模型和数据规模的增大而提升，作者成功训练了一个拥有3.17亿参数的世界模型，'
    '使其能够在多个任务领域、机器人构型和动作空间中完成80项任务。'
)

# 1 背景
add_heading_1(doc, '1 背景')

add_normal(doc,
    '强化学习的目标是通过与环境的交互来学习策略，其形式化描述为马尔可夫决策过程。'
    '本文关注具有连续动作空间的无限时域马尔可夫决策过程，可以将其表示为一个五元组(S, A, T, R, γ)，'
    '其中S为状态空间，A为动作空间，T为状态转移函数，R为与特定任务关联的奖励函数，γ为折扣因子。'
    '目标是找到一个控制策略π，使得期望折扣累积奖励（即回报）最大化。'
)

add_normal(doc,
    '模型预测控制是一种通用的基于模型的控制框架，它通过优化有限长度的动作序列来最大化时间范围H内的回报。'
    '具体来说，模型预测控制求解如下优化问题：在给定当前状态的情况下，找到能使期望回报最大化的动作序列。'
    '候选轨迹的回报通过学习到的模型进行模拟估计。因此，通过该优化问题得到的策略在时间上是局部最优的，'
    '并不能保证（也不太可能）是上述一般强化学习问题的全局最优解。'
    'TD-MPC2通过用学习到的终端值函数对超出规划范围H的回报估计进行自举，来弥补局部轨迹优化的这一不足。'
)

# 2 TD-MPC2
add_heading_1(doc, '2 TD-MPC2')

add_normal(doc,
    '本文在TD-MPC的基础上进行改进。TD-MPC是一种基于模型的强化学习算法，'
    '在学习到的隐式世界模型的潜在空间中进行局部轨迹优化（规划）。'
    'TD-MPC2是用于训练大规模多任务世界模型的实用算法。'
    '具体而言，作者对TD-MPC算法提出了一系列改进，这些改进对于实现较强的算法鲁棒性'
    '（即在所有任务中使用相同超参数）以及将世界模型扩展到比之前多300倍的参数量起到了关键作用。'
)

add_heading_2(doc, '2.1 隐式世界模型的学习')

add_normal(doc,
    '使用重建（解码器）目标来学习环境的生成模型虽然能提供丰富的学习信号，但由于需要在长时间范围内准确预测原始的未来观测值'
    '（如图像或本体感知特征），这一问题本身就很困难，而且并不一定能够带来有效的控制效果。'
    'TD-MPC2没有采用显式建模动力学的重建方式，而是致力于学习一个最为实用的模型：'
    '一个能够根据动作序列准确预测结果（回报）的模型。'
)

add_normal(doc,
    '具体来说，TD-MPC2通过联合嵌入预测、奖励预测和时序差分学习的组合，从环境交互中学习一个隐式的、以控制为中心的世界模型，且不需要解码观测值。'
    '这种基于模型的强化学习的新表述方式，是用适中的模型规模来建模大型数据集的关键所在。'
    '学习到的世界模型随后可以在模型预测控制框架下通过局部轨迹优化（规划）来进行决策。'
)

add_normal(doc,
    'TD-MPC2的架构由五个组件组成：编码器、潜在动力学模型、奖励预测器、终端价值函数和策略先验。'
    '编码器将观测值映射为其潜在表示；潜在动力学模型在潜在空间中建模前向动力学；'
    '奖励预测器预测状态转移的奖励；终端价值函数预测折扣累积回报；'
    '策略先验则预测能使价值函数最大化的动作，用于引导基于采样的轨迹优化器。'
)

add_normal(doc,
    '在训练过程中，编码器、潜在动力学、奖励和价值函数这四个组件通过联合优化来最小化一个综合目标函数。'
    '该目标函数包含三项：联合嵌入预测损失、奖励预测损失和价值预测损失。'
    '由于不同任务之间的奖励量级可能差异很大，TD-MPC2将奖励和价值预测表述为对数变换空间中的离散回归问题，'
    '通过最小化交叉熵来优化。'
)

add_normal(doc,
    '策略先验是一个随机最大熵策略，它学习在保持一定探索性的同时最大化价值估计。'
    '在实践中，作者通过移动统计量来自动调节温度参数α，以平衡价值估计和策略熵之间的关系，防止策略过早坍缩。'
)

add_normal(doc,
    '在架构设计方面，TD-MPC2的所有组件均采用多层感知机实现，中间线性层后接层归一化和Mish激活函数。'
    '为了缓解梯度爆炸问题，作者将潜在表示归一化，通过softmax操作将潜在状态投影到固定维度的单纯形中，'
    '作者将这种归一化方案称为SimNorm。SimNorm的一个关键优势是，它自然地使表示偏向稀疏性，而不需要施加硬约束。'
    '实验表明，SimNorm对于TD-MPC2的训练稳定性至关重要。'
    '此外，为了减少由目标网络产生的价值估计偏差，作者学习了一个Q函数集成体，'
    '并在实践中使用5个Q函数，目标值通过随机子采样的两个目标网络的最小值来计算。'
)

add_heading_2(doc, '2.2 结合策略先验的模型预测控制')

add_normal(doc,
    'TD-MPC2通过使用学习到的世界模型进行规划来获得闭环控制策略。'
    '具体来说，该方法利用模型预测控制框架进行局部轨迹优化，采用模型预测路径积分作为无导数优化器，'
    '通过对长度为H的采样动作序列在模型中展开潜在轨迹来进行评估。'
)

add_normal(doc,
    '在每个决策步骤中，算法估计一个时间相关的多元高斯分布的参数，使得期望回报最大化。'
    '该优化问题通过迭代采样动作序列、评估其期望回报、并基于加权平均更新分布参数来求解。'
    '值得注意的是，该估计通过学习到的终端价值函数对超出规划范围H的回报进行自举，从而估计完整的强化学习目标。'
    'TD-MPC2重复这一迭代规划过程固定的次数，然后执行第一个动作。'
    '为了加速规划收敛，一部分动作序列来自策略先验，并且通过将上一步决策的解移位一步来热启动当前规划。'
)

add_heading_2(doc, '2.3 通用TD-MPC2智能体的训练')

add_normal(doc,
    'TD-MPC2在多样化单任务问题中的成功可以归功于上述算法。'
    '然而，学习一个能够在多个任务领域、机器人构型和动作空间中执行多种任务的大型通用TD-MPC2智能体，会面临若干独特的挑战：'
    '如何学习和表示任务语义？如何在不具备特定领域知识的情况下适应多种观测和动作空间？'
    '如何利用学习到的模型进行新任务的少样本学习？'
)

add_normal(doc,
    '在可学习任务嵌入方面，为了在多任务设置中取得成功，智能体需要学习一种能够利用任务相似性的共享表示，'
    '同时在测试时仍然保留区分不同任务的能力。当存在任务或领域知识（如自然语言指令形式）时，任务嵌入可以编码这些信息。'
    '但在一般情况下，作者选择从数据中学习任务嵌入以及隐含的任务关系。'
    'TD-MPC2用一个可学习的、固定维度的任务嵌入来条件化所有五个组件，该嵌入与模型的其他组件一起联合训练。'
    '为了提高训练稳定性，作者将任务嵌入的L2范数约束在1以内，这也使得实验中的任务嵌入在语义上更加连贯。'
)

add_normal(doc,
    '在动作掩码方面，TD-MPC2学习执行具有不同观测和动作空间的任务，而不依赖任何领域知识。'
    '为此，作者将所有模型输入和输出零填充到各自的最大维度，'
    '并在训练和推理过程中屏蔽策略先验在无效动作维度上的预测。'
    '这确保了无效维度上的预测误差不会影响价值目标的估计，'
    '也防止了策略先验在动作空间较小的任务上错误地膨胀其熵值。'
    '同样，在规划过程中也仅沿有效维度采样动作。'
)

# 3 实验
add_heading_1(doc, '3 实验')

add_normal(doc,
    '作者在总共104个不同的连续控制任务上评估了TD-MPC2，涵盖4个任务领域：DMControl、Meta-World、ManiSkill2和MyoSuite。'
    '这些任务包括高维状态和动作空间（最高达39维）、稀疏奖励、多物体操作、生理精确的肌肉骨骼运动控制、'
    '复杂运动（如四足和双足构型），并覆盖了广泛的任务难度。此外还包括10个具有视觉观测的DMControl任务。'
)

add_normal(doc,
    '实验主要围绕三个核心问题展开：'
    '第一，TD-MPC2与当前最先进的无模型方法（SAC）和基于模型的方法（DreamerV3、TD-MPC）相比表现如何？'
    '第二，TD-MPC2的算法创新能否带来随模型和数据规模增大而持续提升的智能体能力？单一智能体能否学习在多个任务领域、'
    '机器人构型和动作空间中执行多样化技能？'
    '第三，TD-MPC2中引入的具体设计选择如何影响下游任务性能？'
)

add_normal(doc,
    '基线方法包括：软演员评论家算法（SAC），一种基于最大熵强化学习的无模型演员评论家算法；'
    'DreamerV3，一种基于模型的方法，通过学习到的环境生成模型的展开来优化无模型策略；'
    '以及原始版本的TD-MPC。此外，作者还与当前最先进的视觉强化学习方法进行了比较。'
    'SAC和TD-MPC使用针对特定任务调节的超参数，而TD-MPC2在所有任务中使用相同的超参数。'
)

add_heading_2(doc, '3.1 实验结果')

add_normal(doc,
    '与现有方法的比较。作者首先在104个不同任务的在线强化学习设置中，将TD-MPC2的数据效率与一组强基线进行了比较。'
    '结果表明，TD-MPC2在所有任务领域中均优于已有方法。'
    '在高维运动任务（如双足行走和四足奔跑）和多物体操作任务上，TD-MPC2以较大优势超越了基线方法，'
    '尽管它在所有任务中使用了相同的超参数。'
    '值得注意的是，原始TD-MPC有时会因梯度爆炸而发散，而TD-MPC2则保持了稳定。'
)

add_normal(doc,
    '大规模多任务世界模型。为了展示所提改进能够促进世界模型的扩展，作者评估了5个从100万到3.17亿参数的多任务模型'
    '在80个不同任务上的性能。这些任务跨越多个任务领域，在目标、机器人构型和动作空间上差异很大。'
    '模型在一个包含5.45亿条转移的数据集上训练，该数据集来自240个单任务TD-MPC2智能体的经验回放缓冲区，'
    '因此包含了从随机策略到专家策略的多种行为。'
)

add_normal(doc,
    '结果表明，智能体的能力在两个任务集上都随模型规模的增大而持续提升。'
    '对于最大的模型（3.17亿参数），性能在两个数据集上都没有出现饱和的迹象，因此可以预期结果会继续改善。'
    '归一化分数似乎与模型参数量的对数呈线性关系。'
    '3.17亿参数的模型可以在有限的计算资源下完成训练。'
)

add_normal(doc,
    '为了更好地理解多任务模型学习为何成功，作者探索了TD-MPC2学习到的任务嵌入。'
    '有趣的是，语义上相似的任务（如开门和关门）在学习到的任务嵌入空间中距离较近。'
    '然而，嵌入的相似性似乎更多地与任务动力学（机器人构型、物体）而非目标（行走、奔跑）对齐。'
    '这在直觉上是合理的，因为动力学与控制是紧密耦合的。'
)

add_normal(doc,
    '少样本学习。虽然本文主要关注世界模型的扩展和鲁棒性，但作者也探索了将预训练世界模型微调用于未见任务少样本学习的效果。'
    '具体来说，作者在70个任务上预训练了一个1900万参数的TD-MPC2智能体，'
    '然后通过在线强化学习将完整模型微调到10个保留任务中的每一个。'
    '结果表明，在低数据量条件下（2万环境步），TD-MPC2相比从零开始学习提升了2倍。'
)

add_normal(doc,
    '消融实验。作者对TD-MPC2的大部分设计选择进行了消融，包括动作器的选择、各种归一化技术、回归目标和Q函数的数量。'
    '结果表明，所提出的所有改进都对TD-MPC2在单任务和多任务强化学习中的鲁棒性和性能做出了有意义的贡献。'
    '有趣的是，各设计选择的相对重要性在两种设置中是一致的。'
)

add_normal(doc,
    '视觉强化学习。本文主要关注具有本体感知状态观测的高维连续控制任务。'
    '然而，TD-MPC2也可以直接应用于具有其他输入模态的任务。'
    '作者将TD-MPC2的编码器替换为一个浅层卷积编码器，并在10个不同难度的DMControl任务上与当前最先进的视觉强化学习方法进行了比较。'
    '结果表明，TD-MPC2在不改变超参数的情况下，与两个最佳基线（DrQ-v2和DreamerV3）表现相当。'
)

# 4 经验、机遇与风险
add_heading_1(doc, '4 经验、机遇与风险')

add_normal(doc,
    '经验。从历史上看，强化学习算法对架构、超参数、任务特性甚至随机种子都非常敏感，且缺乏调节算法的系统性方法。'
    '因此，深度强化学习的成功应用通常需要拥有大量计算资源的大型专家团队。'
    'TD-MPC2以及若干其他同期的强化学习方法，致力于通过提高现有开源算法的鲁棒性来降低强化学习的使用门槛，'
    '使资源较少的小型学术团队、实践者和个人也能使用。作者坚信，提高算法鲁棒性将继续对这一领域产生深远影响。'
)

add_normal(doc,
    '从TD-MPC2的开发中得到的一个关键经验是，目前还没有发现一种算法能够在各种任务上都表现优异。'
    '例如，DreamerV3在具有离散动作空间的挑战性任务（如Atari游戏和Minecraft）上取得了很好的结果，'
    '但TD-MPC2在困难的连续控制任务上产生了明显更好的结果。与此同时，将TD-MPC2扩展到离散动作空间仍然是一个开放问题。'
)

add_normal(doc,
    '机遇。本文的扩展结果展示了一条基于模型的强化学习路径，其中大规模多任务世界模型可以作为通用世界模型使用。'
    '作者设想了这样一个未来：隐式世界模型可以零样本地用于在已知机器人构型上执行多样化任务，'
    '通过微调快速适应新机器人构型上的任务，并与现有的视觉语言模型结合来执行更高层次的认知任务。'
    '本文的结果是有前景的，但这种程度的泛化可能需要比目前可用的任务多几个数量级的数据。'
)

add_normal(doc,
    '最后，虽然TD-MPC2依赖奖励来进行任务学习，但采用广义的奖励概念（即仅作为任务完成度的度量）是有益的。'
    '这类度量在现实中已经存在，例如成功标签、人类偏好或干预，'
    '或者当前观测与目标之间在已有学习表示中的嵌入距离。然而，利用这类奖励进行大规模预训练仍是一个开放问题。'
)

add_normal(doc,
    '风险。虽然通用世界模型的潜力令人振奋，但仍存在若干挑战：'
    '任务奖励的错误设定可能导致难以预料的不良结果；'
    '在没有额外安全检查的情况下，将物理机器人的无约束自主权交给学习到的模型可能导致灾难性故障；'
    '某些应用的数据获取成本可能过高，导致能力集中在少数团队手中。'
    '缓解这些挑战中的每一项都需要新的研究创新。'
)

# 5 相关工作
add_heading_1(doc, '5 相关工作')

add_normal(doc,
    '多项已有工作致力于构建对超参数、架构以及任务和数据变化具有鲁棒性的强化学习算法。'
    '例如，双Q学习、RED-Q、SVEA和SR-SPR通过调整时序差分目标估计中的偏差方差权衡来提高Q学习算法的稳定性；'
    'C51和DreamerV3通过在变换空间中进行离散回归来提高对奖励量级的鲁棒性；'
    '无模型算法DrQ和DrQ-v2分别通过数据增强和其他重要实现细节来提高训练稳定性和探索效率。'
    '然而，上述所有工作都严格专注于改善单任务在线强化学习中的数据效率和鲁棒性。'
)

add_normal(doc,
    '现有研究神经架构扩展用于决策的文献通常假设可以访问大量近似专家演示数据集用于行为克隆。'
    'Gato通过在海量专家演示数据集上训练基于Transformer的大型序列模型来学习跨领域执行任务，'
    'RT-1同样通过在人类遥操作收集的大型数据集上训练来学习用于物体操作的序列模型。'
    '虽然这些工作的实验结果令人印象深刻，但对大型演示数据集的假设在实际中并不现实。'
    '此外，当前的序列模型依赖于动作空间的离散化（词元化），这使得扩展到高维连续控制任务变得困难。'
)

add_normal(doc,
    '最近，研究者探索了扩展强化学习算法作为上述挑战的解决方案。'
    '例如，VPT通过在大型人类游戏数据集上预训练行为克隆策略，然后用强化学习微调策略来学习玩Minecraft。'
    'DreamerV3成功地扩展了其世界模型的参数规模，并展示了更大的模型在在线强化学习设置中通常具有更高的数据效率，'
    '但并未考虑多任务强化学习。'
)

add_normal(doc,
    '与上述工作不同，TD-MPC2是首个在大规模多任务连续控制问题中展示可扩展性的基于模型的强化学习算法。'
    'TD-MPC2的独特之处在于，它能够在不依赖领域知识的情况下，用单一模型和单一超参数集处理多种机器人构型和动作空间。'
)

# 6 结论
add_heading_1(doc, '6 结论')

add_normal(doc,
    '本文提出了TD-MPC2，一种基于模型的强化学习算法，用于学习可在大规模多任务连续控制问题中使用的通用世界模型。'
    'TD-MPC2在104个不同任务上持续优于现有方法，且只需一组超参数。'
    '研究展示了智能体能力随模型和数据规模增大而提升的路径，'
    '并成功训练了一个3.17亿参数的世界模型来跨多个任务领域、机器人构型和动作空间执行80项任务。'
    '作者公开发布了300多个模型检查点、数据集以及训练和评估代码，以支持开源科学研究。'
)

# ===== 参考文献 =====
add_heading_1(doc, '参考文献')

refs = [
    '[1] Ba J L, Kiros J R, Hinton G E. Layer normalization[J]. Advances in Neural Information Processing Systems, 2016.',
    '[2] Baker B, Akkaya I, Zhokov P, et al. Video pretraining (VPT): Learning to act by watching unlabeled online videos[J]. Advances in Neural Information Processing Systems, 2022, 35: 24639-24654.',
    '[3] Bellemare M G, Dabney W, Munos R. A distributional perspective on reinforcement learning[C]//International Conference on Machine Learning. PMLR, 2017: 449-458.',
    '[4] Bellman R. A markovian decision process[J]. Indiana Univ. Math. J., 1957, 6: 679-684.',
    '[5] Brohan A, Brown N, Carbajal J, et al. RT-2: Vision-language-action models transfer web knowledge to robotic control[J]. arXiv preprint arXiv:2307.15818, 2023.',
    '[6] Brown T, Mann B, Ryder N, et al. Language models are few-shot learners[J]. Advances in Neural Information Processing Systems, 2020, 33: 1877-1901.',
    '[7] Caggiano V, Wang H, Durandau G, et al. MyoSuite - A contact-rich simulation suite for musculoskeletal motor control[J]. 2022.',
    '[8] Calli B, Singh A, Walsman A, et al. The YCB object and model set: Towards common benchmarks for manipulation research[C]//2015 International Conference on Advanced Robotics. IEEE, 2015: 510-517.',
    '[9] Chen X, Wang C, Zhou Z, et al. Randomized ensembled double Q-learning: Learning fast without a model[C]//International Conference on Learning Representations, 2021.',
    '[10] D\'Oro P, Schwarzer M, Nikishin E, et al. Sample-efficient reinforcement learning by breaking the replay ratio barrier[C]//The Eleventh International Conference on Learning Representations, 2023.',
    '[11] Driess D, Xia F, Sajjadi M S M, et al. PaLM-E: An embodied multimodal language model[J]. arXiv preprint arXiv:2303.03378, 2023.',
    '[12] Feng Y, Hansen N, Xiong Z, et al. Finetuning offline world models in the real world[C]//Conference on Robot Learning, 2023.',
    '[13] Grill J B, Strub F, Altch\'e F, et al. Bootstrap your own latent: A new approach to self-supervised learning[J]. Advances in Neural Information Processing Systems, 2020.',
    '[14] Gu J, Xiang F, Li X, et al. ManiSkill2: A unified benchmark for generalizable manipulation skills[C]//International Conference on Learning Representations, 2023.',
    '[15] Haarnoja T, Zhou A, Hartikainen K, et al. Soft actor-critic algorithms and applications[J]. arXiv preprint arXiv:1812.05905, 2018.',
    '[16] Hafner D, Pasukonis J, Ba J, et al. Mastering diverse domains through world models[J]. arXiv preprint arXiv:2301.04104, 2023.',
    '[17] Hansen N, Su H, Wang X. Stabilizing deep Q-learning with convnets and vision transformers under data augmentation[C]//Annual Conference on Neural Information Processing Systems, 2021.',
    '[18] Hansen N, Wang X, Su H. Temporal difference learning for model predictive control[C]//ICML, 2022.',
    '[19] Hansen N, Yuan Z, Ze Y, et al. On pre-training for visuo-motor control: Revisiting a learning-from-scratch baseline[J]. 2023.',
    '[20] Tassa Y, Doron Y, Muldal A, et al. DeepMind control suite[J]. 2018.',
    '[21] Yu T, Quillen D, He Z, et al. Meta-world: A benchmark and evaluation for multi-task and meta reinforcement learning[C]//Conference on Robot Learning. PMLR, 2020: 1094-1100.',
    '[22] Yarats D, Zhang A, Kostrikov I, et al. Improving sample efficiency in model-free reinforcement learning from images[C]//Proceedings of the AAAI Conference on Artificial Intelligence, 2021.',
    '[23] Zhu Z, Yuan Z, Zhou J, et al. On the effectiveness of pretrained representations for visuo-motor learning[J]. 2023.',
]

for ref in refs:
    p = doc.add_paragraph()
    run = p.add_run(ref)
    set_font(run, '宋体', 'Times New Roman', 9)
    p.paragraph_format.line_spacing = 1.5

# ===== 英文版本 =====
doc.add_page_break()
add_title(doc, 'TD-MPC2: Scalable, Robust World Models for Continuous Control')
add_author(doc, 'Nicklas Hansen    Hao Su    Xiaolong Wang')
add_affiliation(doc, '(University of California San Diego, San Diego, CA 92093, USA)')

add_abstract(doc,
    'TD-MPC is a model-based reinforcement learning (RL) algorithm that performs local trajectory optimization in the latent space '
    'of a learned implicit (decoder-free) world model. In this work, we present TD-MPC2: a series of improvements upon the TD-MPC algorithm. '
    'We demonstrate that TD-MPC2 improves significantly over baselines across 104 online RL tasks spanning 4 diverse task domains, '
    'achieving consistently strong results with a single set of hyperparameters. We further show that agent capabilities increase with model '
    'and data size, and successfully train a single 317M parameter agent to perform 80 tasks across multiple task domains, embodiments, '
    'and action spaces. We conclude with an account of lessons, opportunities, and risks associated with large TD-MPC2 agents.'
)

add_keywords(doc,
    'model-based reinforcement learning; world models; model predictive control; continuous control; multi-task learning'
)

# ===== 保存 =====
output_path = r'E:\Code\Embed\作业\智能机器人控制\TD-MPC2-中文翻译-郑州大学工学报格式.docx'
doc.save(output_path)
print(f'Translation saved to: {output_path}')
